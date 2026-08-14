"""
docx_export.py
==============
Builds a downloadable Word document (.docx) from the session's questions and
answers. The interesting part is the MATH: the answers contain LaTeX, and Word
has no built-in way to take "$\\int_0^1 x\\,dx$" and show a real equation.

Strategy (self-contained, no LaTeX install required):
    For every piece of math we find, we render it to a small transparent PNG
    with matplotlib's "mathtext" engine, then drop that image into the document.
    Plain text becomes normal Word runs; **bold** and headings are honored.

matplotlib's mathtext supports a large, common subset of LaTeX (fractions,
integrals, sums, Greek letters, sub/superscripts, etc.). If a particular
expression is too exotic to render, we fall back to printing the raw LaTeX as
text so nothing is ever silently lost.
"""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO

import matplotlib
matplotlib.use("Agg")            # headless backend — no display needed on a server
import matplotlib.pyplot as plt
from PIL import Image

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
MATH_DPI = 200                   # resolution of rendered equations
INLINE_MATH_HEIGHT_PT = 12       # on-screen height of inline equations, in points
DISPLAY_MAX_WIDTH_IN = 6.0       # cap width of standalone equations, in inches
NAVY = RGBColor(0x13, 0x33, 0x5C)  # heading color (matches your usual palette)

# A small cache so the same equation isn't rendered twice in one document.
_render_cache: dict[str, bytes] = {}


# ---------------------------------------------------------------------------
# Math rendering: LaTeX string  ->  PNG bytes
# ---------------------------------------------------------------------------

def _sanitize_for_mathtext(latex: str) -> str:
    """
    matplotlib's mathtext supports a big subset of LaTeX but not everything.
    Rewrite a few common commands it lacks into equivalents it understands.
    (On-screen rendering uses KaTeX, which handles these natively — this only
    helps the Word-export path.)
    """
    # \xrightarrow{d}  ->  \rightarrow^{d}  (keeps the label as a superscript)
    latex = re.sub(r"\\xrightarrow\s*\{(.*?)\}", r"\\rightarrow^{\1}", latex)
    latex = re.sub(r"\\xleftarrow\s*\{(.*?)\}", r"\\leftarrow^{\1}", latex)
    # \text{...} / \operatorname{...}  ->  \mathrm{...}
    latex = re.sub(r"\\text\s*\{", r"\\mathrm{", latex)
    latex = re.sub(r"\\operatorname\s*\{", r"\\mathrm{", latex)
    return latex


def _render_latex_png(latex: str, fontsize: int = 16) -> bytes | None:
    """
    Render a LaTeX snippet (WITHOUT surrounding $ signs) to PNG bytes.
    Tries the expression as-is, then a sanitized version; returns None only if
    both fail, so the caller can fall back to printing raw text.
    """
    key = f"{fontsize}:{latex}"
    if key in _render_cache:
        return _render_cache[key]

    # Try the original first, then a sanitized rewrite of unsupported commands.
    for candidate in (latex, _sanitize_for_mathtext(latex)):
        try:
            fig = plt.figure()
            fig.text(0, 0, f"${candidate}$", fontsize=fontsize)  # mathtext needs $...$
            buf = BytesIO()
            # bbox_inches="tight" crops the figure tightly around the equation.
            fig.savefig(buf, format="png", dpi=MATH_DPI, transparent=True,
                        bbox_inches="tight", pad_inches=0.05)
            plt.close(fig)
            buf.seek(0)
            data = buf.getvalue()
            _render_cache[key] = data
            return data
        except Exception:
            plt.close("all")
            continue
    return None  # both attempts failed -> caller prints raw text


def _png_aspect(png: bytes) -> float:
    """Return width/height of a PNG so we can size it while keeping proportions."""
    with Image.open(BytesIO(png)) as im:
        w, h = im.size
    return w / h if h else 1.0


# ---------------------------------------------------------------------------
# Delimiter normalization + tokenizing a line into text / bold / math pieces
# ---------------------------------------------------------------------------

def _normalize_delimiters(text: str) -> str:
    """Convert \\[..\\] and \\(..\\) into $$..$$ and $..$ so we only parse dollars."""
    text = text.replace(r"\[", "$$").replace(r"\]", "$$")
    text = text.replace(r"\(", "$").replace(r"\)", "$")
    return text


# Matches (in priority order): display math, inline math, **bold**.
_INLINE_TOKEN_RE = re.compile(r"(\$\$.+?\$\$|\$.+?\$|\*\*.+?\*\*)", re.DOTALL)
# A whole line that is exactly one display equation, e.g.  $$ ... $$
_DISPLAY_LINE_RE = re.compile(r"^\$\$(.+?)\$\$$", re.DOTALL)


def _add_inline_content(paragraph, text: str) -> None:
    """
    Add a line's content to an existing paragraph, handling inline math and
    bold. Inline equations are inserted as small images; everything else is text.
    """
    for token in _INLINE_TOKEN_RE.split(text):
        if not token:
            continue
        if token.startswith("$$") and token.endswith("$$"):
            latex = token[2:-2].strip()
            _add_inline_math(paragraph, latex)
        elif token.startswith("$") and token.endswith("$"):
            latex = token[1:-1].strip()
            _add_inline_math(paragraph, latex)
        elif token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            paragraph.add_run(token)


def _add_inline_math(paragraph, latex: str) -> None:
    """Insert an inline equation image into a paragraph (or fall back to text)."""
    png = _render_latex_png(latex)
    if png is None:
        paragraph.add_run(latex)  # fallback: show the LaTeX source as plain text
        return
    run = paragraph.add_run()
    # Sizing by height keeps the equation aligned with the surrounding text.
    run.add_picture(BytesIO(png), height=Pt(INLINE_MATH_HEIGHT_PT))


def _add_display_math(doc: Document, latex: str) -> None:
    """Insert a standalone, centered equation image on its own line."""
    png = _render_latex_png(latex, fontsize=18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if png is None:
        p.add_run(latex)  # fallback: show the LaTeX source as plain text
        return
    aspect = _png_aspect(png)
    height_in = 0.32                          # base height for a display equation
    width_in = min(DISPLAY_MAX_WIDTH_IN, height_in * aspect)
    p.add_run().add_picture(BytesIO(png), width=Inches(width_in))


# ---------------------------------------------------------------------------
# Render one answer (markdown + LaTeX) into the document
# ---------------------------------------------------------------------------

def _render_answer(doc: Document, answer: str) -> None:
    """Walk the answer line by line, mapping markdown structure to Word."""
    answer = _normalize_delimiters(answer)

    for raw_line in answer.split("\n"):
        line = raw_line.strip()
        if not line:
            continue  # blank line -> just a visual gap; skip

        # Standalone display equation on its own line.
        m = _DISPLAY_LINE_RE.match(line)
        if m:
            _add_display_math(doc, m.group(1).strip())
            continue

        # Markdown headings (### / ## / #) -> Word heading levels.
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue

        # Bullet list items.
        if line.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_content(p, line[2:].strip())
            continue

        # Ordinary paragraph (may contain inline math / bold).
        p = doc.add_paragraph()
        _add_inline_content(p, line)


# ---------------------------------------------------------------------------
# PUBLIC ENTRY POINT — build the whole document and return its bytes
# ---------------------------------------------------------------------------

def build_qa_document(qa_pairs: list[dict]) -> bytes:
    """
    Build a .docx from a list of Q&A items and return it as bytes (ready for a
    Streamlit download_button).

    Each item is a dict:
        {"question": str, "answer": str, "sources": str | None}
    """
    _render_cache.clear()
    doc = Document()

    # Base body font.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Cover heading.
    title = doc.add_heading("Math & Statistics Q&A — Session Notes", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY
    doc.add_paragraph(f"Generated {date.today().isoformat()}")

    # One block per question.
    for i, item in enumerate(qa_pairs, start=1):
        question = item.get("question", "").strip()
        answer = item.get("answer", "").strip()
        sources = item.get("sources")

        # Question as a heading (questions are usually plain text).
        doc.add_heading(f"Q{i}. {question}", level=2)

        # Answer body (handles structure + equations).
        _render_answer(doc, answer)

        # Optional sources caption in small grey italics.
        if sources:
            cap = doc.add_paragraph()
            run = cap.add_run(sources)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x5B, 0x6B, 0x7E)

        # Spacer between items (except after the last one).
        if i < len(qa_pairs):
            doc.add_paragraph()

    # Save to an in-memory buffer and hand back the raw bytes.
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
